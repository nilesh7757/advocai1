from rest_framework.decorators import api_view, parser_classes, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from django.http import FileResponse
from django.core.mail import send_mail
from django.conf import settings
import re
import markdown
from xhtml2pdf import pisa
from io import BytesIO
from rest_framework import status
from rest_framework.parsers import MultiPartParser, FormParser
import cloudinary.uploader
from documents.mongo_client import get_conversation_by_id
from utils.models import SupportQuery
from utils.permissions import IsAdminUser


@api_view(['POST'])
def download_pdf(request):
    """
    API endpoint to download a legal document as PDF.
    """
    document_content = request.data.get('document_content')
    if not document_content:
        return Response({'error': 'Document content is required'}, status=400)

    try:
        pdf_file = _generate_pdf_from_markdown(document_content)
        response = FileResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="legal_document.pdf"'
        return response
    except Exception as e:
        return Response({'error': f'Error generating PDF: {e}'}, status=500)

def _generate_pdf_from_markdown(markdown_content):
    """
    Helper function to convert markdown string to a PDF file response.
    This function is used by the download_pdf view.
    """
    print(f"Markdown content received: {markdown_content[:500]}...") # Log first 500 chars
    html_content = markdown.markdown(markdown_content)
    print(f"HTML content generated: {html_content[:500]}...") # Log first 500 chars

    pdf_style_css = """
        @page {
            size: a4 portrait;
            margin: 1.2cm;
        }
        body {
            font-family: "Times New Roman", Times, serif;
            font-size: 11pt;
            line-height: 1.3;
            color: #000000;
        }
        h1, h2, h3, h4, h5, h6 {
            font-family: "Times New Roman", Times, serif;
            font-weight: bold;
            color: #000000;
            margin-top: 1.2em;
            margin-bottom: 0.6em;
            line-height: 1.15;
        }
        h1 {
            font-size: 16pt;
            text-align: center;
            text-transform: uppercase;
            margin-bottom: 1.5em;
        }
        h2 {
            font-size: 14pt;
            text-transform: uppercase;
            border-bottom: 1px solid #000000;
            padding-bottom: 0.2em;
        }
        h3 {
            font-size: 12pt;
            font-weight: bold;
            text-decoration: underline;
        }
        p {
            margin-bottom: 0.8em;
            text-align: justify;
            text-indent: 1.25cm; /* Indent first line of paragraphs */
        }
        /* Don't indent first paragraph after a heading */
        h1 + p, h2 + p, h3 + p, h4 + p, h5 + p, h6 + p {
            text-indent: 0;
        }
        ul, ol {
            margin-bottom: 0.8em;
            padding-left: 1.5cm;
        }
        li {
            margin-bottom: 0.3em;
            text-align: justify;
        }
        strong, b {
            font-weight: bold;
        }
        em, i {
            font-style: italic;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 1em;
            border: 1px solid #333333;
        }
        th, td {
            border: 1px solid #333333;
            padding: 6px;
            text-align: left;
            vertical-align: top;
        }
        th {
            background-color: #e0e0e0;
            font-weight: bold;
        }
        hr {
            width: 250px;
            margin-left: 0;
            border: 0.5px solid #000;
        }
        /* Signature sizing and spacing */
        img[alt~="signature"][alt~="landlord"] {
            display: block;
            width: 180px;
            height: 80px;
            object-fit: contain;
            margin-top: 8mm;   /* place below landlord text */
            margin-bottom: 0;
        }
        img[alt~="signature"][alt~="tenant"] {
            display: block;
            width: 180px;
            height: 80px;
            object-fit: contain;
            margin-top: 0;
            margin-bottom: 8mm; /* place above tenant text */
        }
        /* Remove header and footer for a more traditional look */
    """

    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Legal Document</title>
        <meta charset=\"utf-8\">
        <style>{pdf_style_css}</style>
    </head>
    <body>{html_content}</body>
    </html>
    """

    result_file = BytesIO()
    pisa_status = pisa.CreatePDF(full_html, dest=result_file)

    if pisa_status.err:
        raise Exception(f'PDF generation error: {pisa_status.err}')

    result_file.seek(0)
    return result_file

@api_view(['POST'])
@parser_classes([MultiPartParser, FormParser])
def upload_signature(request):
    """
    Accepts an image upload and returns its accessible URL for embedding in markdown.
    """
    file_obj = request.FILES.get('signature') or request.FILES.get('file')
    if not file_obj:
        return Response({'error': 'No file uploaded. Use form field name "signature".'}, status=400)

    try:
        upload_result = cloudinary.uploader.upload(file_obj)
        print(f"Cloudinary upload result: {upload_result}") # Add this line for debugging
        return Response({'url': upload_result['secure_url']}, status=201)
    except Exception as e:
        return Response({'error': str(e)}, status=500)


@api_view(['GET'])
def download_latest_conversation_pdf(request, pk):
    """
    Downloads the latest document content from a conversation as a PDF.
    """
    conversation = get_conversation_by_id(pk)
    if not conversation or 'document_versions' not in conversation or not conversation['document_versions']:
        return Response({'error': 'No document content found for this conversation.'}, status=status.HTTP_404_NOT_FOUND)

    try:
        # Get the content of the latest version
        latest_version_content = conversation['document_versions'][-1]['content']
        pdf_file = _generate_pdf_from_markdown(latest_version_content)
        
        response = FileResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{conversation.get("title", "legal_document")}.pdf"'
        return response
    except Exception as e:
        return Response({'error': f'Error generating PDF: {e}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
def download_version_pdf(request, pk, version_number):
    """
    Downloads a specific document version from a conversation as a PDF.
    """
    try:
        conversation = get_conversation_by_id(pk)
        if not conversation or 'document_versions' not in conversation or not conversation['document_versions']:
            return Response({'error': 'No document versions found for this conversation.'}, status=status.HTTP_404_NOT_FOUND)
        
        version = next((v for v in conversation['document_versions'] if v['version_number'] == version_number), None)
        if not version:
            return Response({'error': 'Version content not found'}, status=status.HTTP_404_NOT_FOUND)

        pdf_file = _generate_pdf_from_markdown(version['content'])
        filename = f"{conversation.get('title', 'legal_document')}_v{version_number}.pdf"
        
        response = FileResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        print(f"Error in download_version_pdf: {e}")
        return Response({'error': f'Error generating PDF: {e}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def _generate_docx_from_markdown(content):
    from docx import Document
    from html.parser import HTMLParser
    import markdown
    from io import BytesIO

    # Convert markdown to html first (which keeps HTML tags unchanged)
    html_content = markdown.markdown(content)

    doc = Document()

    class DocxHTMLParser(HTMLParser):
        def __init__(self, doc):
            super().__init__()
            self.doc = doc
            self.current_paragraph = None
            self.is_bold = False
            self.is_italic = False
            self.is_underline = False
            self.list_type = None  # 'ul' or 'ol'
            
        def handle_starttag(self, tag, attrs):
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1])
                self.current_paragraph = self.doc.add_heading('', level=level)
            elif tag == 'p':
                self.current_paragraph = self.doc.add_paragraph()
            elif tag == 'li':
                style = 'List Bullet' if self.list_type == 'ul' else 'List Number'
                self.current_paragraph = self.doc.add_paragraph(style=style)
            elif tag == 'ul':
                self.list_type = 'ul'
            elif tag == 'ol':
                self.list_type = 'ol'
            elif tag == 'strong' or tag == 'b':
                self.is_bold = True
            elif tag == 'em' or tag == 'i':
                self.is_italic = True
            elif tag == 'u':
                self.is_underline = True
            elif tag == 'br':
                if self.current_paragraph:
                    self.current_paragraph.add_run('\n')
                else:
                    self.current_paragraph = self.doc.add_paragraph()
                    
        def handle_endtag(self, tag):
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']:
                self.current_paragraph = None
            elif tag in ['ul', 'ol']:
                self.list_type = None
            elif tag == 'strong' or tag == 'b':
                self.is_bold = False
            elif tag == 'em' or tag == 'i':
                self.is_italic = False
            elif tag == 'u':
                self.is_underline = False
                
        def handle_data(self, data):
            if not data.strip() and not self.current_paragraph:
                return
            
            if not self.current_paragraph:
                self.current_paragraph = self.doc.add_paragraph()
                
            run = self.current_paragraph.add_run(data)
            if self.is_bold:
                run.bold = True
            if self.is_italic:
                run.italic = True
            if self.is_underline:
                run.underline = True

    parser = DocxHTMLParser(doc)
    parser.feed(html_content)

    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file


@api_view(['POST'])
def download_docx(request):
    """
    API endpoint to download a legal document as DOCX.
    """
    document_content = request.data.get('document_content')
    if not document_content:
        return Response({'error': 'Document content is required'}, status=400)

    try:
        docx_file = _generate_docx_from_markdown(document_content)
        response = FileResponse(docx_file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="legal_document.docx"'
        return response
    except Exception as e:
        return Response({'error': f'Error generating DOCX: {e}'}, status=500)


@api_view(['GET'])
def download_latest_conversation_docx(request, pk):
    """
    Downloads the latest document content from a conversation as a DOCX file.
    """
    conversation = get_conversation_by_id(pk)
    if not conversation or 'document_versions' not in conversation or not conversation['document_versions']:
        return Response({'error': 'No document content found for this conversation.'}, status=status.HTTP_404_NOT_FOUND)

    try:
        latest_version_content = conversation['document_versions'][-1]['content']
        docx_file = _generate_docx_from_markdown(latest_version_content)
        
        response = FileResponse(docx_file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{conversation.get("title", "legal_document")}.docx"'
        return response
    except Exception as e:
        return Response({'error': f'Error generating DOCX: {e}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def download_version_docx(request, pk, version_number):
    """
    Downloads a specific document version from a conversation as a DOCX file.
    """
    try:
        conversation = get_conversation_by_id(pk)
        if not conversation or 'document_versions' not in conversation or not conversation['document_versions']:
            return Response({'error': 'No document versions found for this conversation.'}, status=status.HTTP_404_NOT_FOUND)
        
        version = next((v for v in conversation['document_versions'] if v['version_number'] == version_number), None)
        if not version:
            return Response({'error': 'Version content not found'}, status=status.HTTP_404_NOT_FOUND)

        docx_file = _generate_docx_from_markdown(version['content'])
        filename = f"{conversation.get('title', 'legal_document')}_v{version_number}.docx"
        
        response = FileResponse(docx_file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        print(f"Error in download_version_docx: {e}")
        return Response({'error': f'Error generating DOCX: {e}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([AllowAny])
def contact(request):
    name = (request.data.get('name') or '').strip()
    email = (request.data.get('email') or '').strip()
    subject = (request.data.get('subject') or '').strip()
    message = (request.data.get('message') or '').strip()

    errors = {}
    if not name:
        errors['name'] = 'Name is required.'
    if not email:
        errors['email'] = 'Email is required.'
    elif not re.match(r'^[^\s@]+@[^\s@]+\.[^\s@]+$', email):
        errors['email'] = 'Please enter a valid email address.'
    if not subject:
        errors['subject'] = 'Subject is required.'
    if not message:
        errors['message'] = 'Message is required.'

    if errors:
        return Response(errors, status=status.HTTP_400_BAD_REQUEST)

    email_subject = f'Contact Form: {subject}'
    email_message = f'From: {name} <{email}>\n\n{message}'
    from_email = settings.DEFAULT_FROM_EMAIL
    recipient = getattr(settings, 'SUPPORT_EMAIL', None) or from_email

    try:
        if settings.DEBUG and not getattr(settings, 'EMAIL_HOST_USER', None):
            print('\n' + '=' * 60)
            print(f'[DEV MODE] Contact form submission')
            print(f'From: {name} <{email}>')
            print(f'Subject: {subject}')
            print(f'Message:\n{message}')
            print('=' * 60 + '\n')
        else:
            send_mail(email_subject, email_message, from_email, [recipient])
    except Exception as e:
        print(f'Error sending contact email: {e}')
        if settings.DEBUG:
            print('\n' + '=' * 60)
            print(f'[DEV MODE FALLBACK] Contact form submission')
            print(f'From: {name} <{email}>')
            print(f'Subject: {subject}')
            print(f'Message:\n{message}')
            print('=' * 60 + '\n')
        else:
            return Response({'error': 'Failed to send message. Please try again later.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # Save the SupportQuery
    try:
        user_ref = request.user if request.user and request.user.is_authenticated else None
        SupportQuery(
            user=user_ref,
            name=name,
            email=email,
            subject=subject,
            message=message
        ).save()
    except Exception as e:
        print(f"Error saving SupportQuery: {e}")

    return Response({'message': 'Your message has been sent. We typically respond within 1-2 business days.'}, status=status.HTTP_200_OK)


# ==========================================
# Admin Support Query Endpoints
# ==========================================

def _serialize_support_query(query):
    user_data = None
    if query.user:
        user_data = {
            'id': str(query.user.id),
            'username': query.user.username,
            'email': query.user.email,
        }
    return {
        'id': str(query.id),
        'user': user_data,
        'name': query.name,
        'email': query.email,
        'subject': query.subject,
        'message': query.message,
        'status': query.status,
        'admin_reply': getattr(query, 'admin_reply', '') or '',
        'replied_by': str(query.replied_by.id) if query.replied_by else None,
        'replied_at': query.replied_at.isoformat() if query.replied_at else None,
        'created_at': query.created_at.isoformat() if query.created_at else None,
    }


@api_view(['GET'])
@permission_classes([IsAdminUser])
def admin_list_queries(request):
    """List all SupportQuery records, most recent first, with optional status filter"""
    status_filter = request.query_params.get('status')
    
    queries = SupportQuery.objects.all()
    if status_filter:
        queries = queries.filter(status=status_filter)
        
    queries = queries.order_by('-created_at')
    data = [_serialize_support_query(q) for q in queries]
    return Response(data, status=status.HTTP_200_OK)


@api_view(['PATCH'])
@permission_classes([IsAdminUser])
def admin_reply_query(request, query_id):
    """Reply to a SupportQuery and notify or email the user"""
    reply = request.data.get('reply', '').strip()
    if not reply:
        return Response({'error': 'Reply message is required.'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        query = SupportQuery.objects(id=query_id).first()
    except Exception:
        query = None

    if not query:
        return Response({'error': 'Support query not found.'}, status=status.HTTP_404_NOT_FOUND)

    from datetime import datetime
    query.admin_reply = reply
    query.replied_by = request.user
    query.replied_at = datetime.utcnow()
    query.status = 'answered'
    query.save()

    # If the query is associated with a registered User, send an in-app notification
    if query.user:
        try:
            from authentication.models import Notification
            Notification(
                recipient=query.user,
                sender=request.user,
                notification_type='support_reply',
                document_id='system',
                message=f"Admin replied to your query '{query.subject}': {reply[:100]}..."
            ).save()
        except Exception as e:
            print(f"Error creating in-app notification for query reply: {e}")
    else:
        # If anonymous submission, send an email to the provided email address
        try:
            email_subject = f"Re: {query.subject}"
            email_message = (
                f"Hi {query.name},\n\n"
                f"Thank you for contacting AdvocAI Support. We have replied to your message:\n\n"
                f"\"...{query.message}...\"\n\n"
                f"--- Support Reply ---\n"
                f"{reply}\n\n"
                f"Best regards,\n"
                f"AdvocAI Team"
            )
            from_email = settings.DEFAULT_FROM_EMAIL
            send_mail(email_subject, email_message, from_email, [query.email])
        except Exception as e:
            print(f"Error sending email reply: {e}")

    return Response(_serialize_support_query(query), status=status.HTTP_200_OK)


@api_view(['PATCH'])
@permission_classes([IsAdminUser])
def admin_close_query(request, query_id):
    """Close a SupportQuery without replying"""
    try:
        query = SupportQuery.objects(id=query_id).first()
    except Exception:
        query = None

    if not query:
        return Response({'error': 'Support query not found.'}, status=status.HTTP_404_NOT_FOUND)

    query.status = 'closed'
    query.save()

    return Response(_serialize_support_query(query), status=status.HTTP_200_OK)